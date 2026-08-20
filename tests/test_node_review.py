import json
from pathlib import Path

from docx import Document
from pydantic_ai import Agent
from pydantic_ai.messages import ModelResponse, ToolCallPart
from pydantic_ai.models.function import AgentInfo, FunctionModel

from biaoshu_gen.nodes import review as rv
from biaoshu_gen.schemas import ReviewReport
from biaoshu_gen.state import BidState, run_dir


def _setup(tmp_path: Path, monkeypatch) -> BidState:
    monkeypatch.chdir(tmp_path)
    state = BidState(run_id="run-1")
    d = run_dir(state)
    parse = d / "01_parse"
    parse.mkdir(parents=True)
    (parse / "invalidation.yaml").write_text("items: []\n", encoding="utf-8")
    (parse / "scoring.yaml").write_text("technical_rules: []\n", encoding="utf-8")
    (d / "03_facts.yaml").write_text("schedule: 90 天\n", encoding="utf-8")
    (d / "02_template").mkdir(parents=True, exist_ok=True)
    (d / "02_template" / "template.md").write_text("# 目录\n- 商务部分\n", encoding="utf-8")
    draft = d / "07_draft" / "标书草稿_v1.docx"
    draft.parent.mkdir(parents=True)
    doc = Document()
    doc.add_heading("技术方案", level=1)
    doc.add_paragraph("正文内容。")
    doc.save(draft)
    return state.model_copy(update={"draft_docx_path": str(draft)})


def _factory(report: dict):
    def make(output_type, system_prompt, retries=2):
        async def fn(messages, info: AgentInfo):
            tool_name = info.output_tools[0].name if info.output_tools else "final_result"
            return ModelResponse(parts=[ToolCallPart(tool_name=tool_name, args=json.dumps(report))])
        return Agent(model=FunctionModel(fn), output_type=output_type,
                     system_prompt=system_prompt, retries=retries)
    return make


def test_review_pass_writes_report(tmp_path: Path, monkeypatch):
    state = _setup(tmp_path, monkeypatch)
    monkeypatch.setattr(rv, "make_agent", _factory({
        "passed": True, "aspects": [{"name": "废标项+扣分项", "passed": True, "note": "无触犯"}],
        "issues": []}))
    updates = rv.review_node(state)
    assert updates["review_passed"] is True
    out = Path(updates["review_report_path"])
    text = out.read_text(encoding="utf-8")
    assert "VERDICT: PASS" in text and "废标项+扣分项" in text
    assert "需人工处理" not in text


def test_review_fail_at_cap_appends_manual_note(tmp_path: Path, monkeypatch):
    state = _setup(tmp_path, monkeypatch).model_copy(update={"revision_round": 2})
    monkeypatch.setattr(rv, "make_agent", _factory({
        "passed": False, "aspects": [{"name": "材料齐全性", "passed": False, "note": "缺表"}],
        "issues": ["缺《投标人基本情况表》"]}))
    updates = rv.review_node(state)
    assert updates["review_passed"] is False
    out = Path(updates["review_report_path"])
    text = out.read_text(encoding="utf-8")
    assert "VERDICT: FAIL" in text and "需人工处理" in text


def test_review_fail_below_cap_no_manual_note(tmp_path: Path, monkeypatch):
    state = _setup(tmp_path, monkeypatch).model_copy(update={"revision_round": 0})
    monkeypatch.setattr(rv, "make_agent", _factory({
        "passed": False, "aspects": [], "issues": ["报价未填"]}))
    updates = rv.review_node(state)
    assert updates["review_passed"] is False
    text = Path(updates["review_report_path"]).read_text(encoding="utf-8")
    assert "需人工处理" not in text
