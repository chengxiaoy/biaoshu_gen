"""revise 节点 P0/P1/P2 优化项的单元测试。"""
from pathlib import Path

from biaoshu_gen.nodes.revise import _write_draft_map
from biaoshu_gen.prompts.revise import build_user_prompt


def _tiny_docx(path: Path) -> Path:
    import docx
    d = docx.Document()
    d.add_paragraph("项目名称：测试项目")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "序号"
    t.cell(0, 1).text = "投标报价"
    d.save(str(path))
    return path


def test_write_draft_map(tmp_path: Path):
    draft = _tiny_docx(tmp_path / "draft.docx")
    ws = tmp_path / "ws"
    ws.mkdir()
    m = _write_draft_map(ws, draft)
    assert m == ws / "_map.txt" and m.exists()
    text = m.read_text(encoding="utf-8")
    assert "项目名称：测试项目" in text
    assert "投标报价" in text


def test_write_draft_map_survives_bad_input(tmp_path: Path):
    """地图生成失败不抛异常（agent 可自行 dump）。"""
    assert _write_draft_map(tmp_path, tmp_path / "不存在.docx") is None


def test_prompt_references_map_and_env():
    p = build_user_prompt("标书草稿_v2.docx", 2, current="标书草稿_v1.docx",
                          env="环境说明：python 是 /x/py.exe")
    assert "_map.txt" in p and "grep" in p
    assert "不要自己重新 dump" in p
    assert "标书草稿_v1.docx" in p and "标书草稿_v2.md" in p
    assert "/x/py.exe" in p                      # env 段落在 prompt 开头


def test_prompt_default_env_empty():
    p = build_user_prompt("标书草稿_v2.docx", 2)
    assert not p.startswith("\n")               # env 缺省时无空段残留
    assert "_map.txt" in p
