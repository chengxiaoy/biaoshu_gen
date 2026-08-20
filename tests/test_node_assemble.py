from pathlib import Path

from docx import Document

from biaoshu_gen.nodes import assemble as asm
from biaoshu_gen.schemas import TenderMetadata
from biaoshu_gen.state import BidState, run_dir


def _forms_docx(path: Path) -> None:
    """底稿候选：模板壳 + 已填的投标函（商务/技术仍是空壳）。"""
    d = Document()
    d.add_heading("投标函", level=1)
    d.add_paragraph("公司：测试投标人公司")
    d.add_heading("资格证明文件", level=1)
    d.add_paragraph("营业执照复印件")
    d.add_heading("商务部分", level=1)
    d.add_paragraph("（此处附业绩与团队）")
    d.add_heading("技术部分", level=1)
    d.add_paragraph("（技术方案格式自定）")
    d.save(path)


def _commercial_docx(path: Path) -> None:
    """整本模板副本，其中商务部分已填。"""
    d = Document()
    d.add_heading("投标函", level=1)
    d.add_paragraph("公司：测试投标人公司")
    d.add_heading("资格证明文件", level=1)
    d.add_paragraph("营业执照复印件")
    d.add_heading("商务部分", level=1)
    d.add_paragraph("业绩：智慧城市监测平台合同")
    d.add_heading("技术部分", level=1)
    d.add_paragraph("（技术方案格式自定）")
    d.save(path)


def _deviation_docx(path: Path) -> None:
    """整本模板副本 + 偏离表（底稿没有的区间）。"""
    d = Document()
    d.add_heading("投标函", level=1)
    d.add_paragraph("公司：测试投标人公司")
    d.add_heading("偏离表", level=1)
    d.add_paragraph("偏离说明：全部无偏离")
    d.save(path)


def _state(tmp_path: Path, monkeypatch, version: int = 0, **paths) -> BidState:
    monkeypatch.chdir(tmp_path)
    d = run_dir(BidState(run_id="run-1"))
    body = d / "05_body"
    body.mkdir(parents=True, exist_ok=True)
    (body / "body.md").write_text("# 1 总体思路\n\n总体思路内容。", encoding="utf-8")
    for name in ("forms", "commercial", "deviation"):
        p = d / "06_fill" / name / f"{name}.docx"
        p.parent.mkdir(parents=True, exist_ok=True)
        maker = {"forms": _forms_docx, "commercial": _commercial_docx,
                 "deviation": _deviation_docx}[name]
        if paths.get(name, True):
            maker(p)
    return BidState(
        run_id="run-1",
        metadata=TenderMetadata(project_name="演示项目"),
        body_md_path=str(body / "body.md"),
        forms_docx_path=str(d / "06_fill/forms/forms.docx") if paths.get("forms", True) else "",
        deviation_docx_path=str(d / "06_fill/deviation/deviation.docx") if paths.get("deviation") else "",
        commercial_docx_path=str(d / "06_fill/commercial/commercial.docx") if paths.get("commercial") else "",
        draft_version=version,
    )


def test_assemble_replaces_anchored_sections_in_template_order(tmp_path: Path, monkeypatch):
    """商务/技术按锚标题区间替换（非追加）：壳不重复、顺序跟模板、正文进锚点。"""
    state = _state(tmp_path, monkeypatch, forms=True, commercial=True, deviation=False)
    updates = asm.assemble_node(state)
    doc = Document(updates["draft_docx_path"])
    texts = [p.text for p in doc.paragraphs]

    # 各锚标题只出现一次（无重复壳）
    for h in ("投标函", "商务部分", "技术部分"):
        assert texts.count(h) == 1, (h, texts)
    # 填充内容在位
    assert "公司：测试投标人公司" in texts            # forms 底稿
    assert "业绩：智慧城市监测平台合同" in texts        # commercial 商务区间（替换）
    # 技术部分空壳被正文替换
    assert "（技术方案格式自定）" not in texts
    assert "总体思路内容。" in texts
    # 顺序跟模板：商务内容 < 技术部分标题 < 正文内容
    assert texts.index("业绩：智慧城市监测平台合同") < texts.index("技术部分")
    assert texts.index("技术部分") < texts.index("总体思路内容。")


def test_assemble_appends_only_missing_deviation_range(tmp_path: Path, monkeypatch):
    """底稿无偏离表区间 -> 仅追加填充文档中的偏离表区间，不整本拼接。"""
    state = _state(tmp_path, monkeypatch, forms=True, commercial=False, deviation=True)
    updates = asm.assemble_node(state)
    doc = Document(updates["draft_docx_path"])
    texts = [p.text for p in doc.paragraphs]

    assert texts.count("偏离表") == 1
    assert "偏离说明：全部无偏离" in texts
    assert texts.count("投标函") == 1               # deviation 整本未拼接进来


def test_assemble_fallback_without_heading_styles(tmp_path: Path, monkeypatch):
    """无标题样式的底稿：正文尾部追加 + commercial 兜底去重追加，不崩溃。"""
    monkeypatch.chdir(tmp_path)
    d = run_dir(BidState(run_id="run-1"))
    body = d / "05_body"
    body.mkdir(parents=True, exist_ok=True)
    (body / "body.md").write_text("# 1 总体思路\n\n总体思路内容。", encoding="utf-8")

    plain = Document()                                # 无 Heading 样式
    plain.add_paragraph("封面页")
    plain.add_paragraph("商务部分")
    plain.add_paragraph("（此处填写）")
    forms_p = d / "06_fill" / "forms" / "forms.docx"
    forms_p.parent.mkdir(parents=True, exist_ok=True)
    plain.save(forms_p)

    comm = Document()                                 # 整本无标题 + 商务已填
    comm.add_paragraph("封面页")
    comm.add_paragraph("商务部分")
    comm.add_paragraph("（此处填写）")
    comm.add_paragraph("业绩：合同一份")
    comm_p = d / "06_fill" / "commercial" / "commercial.docx"
    comm_p.parent.mkdir(parents=True, exist_ok=True)
    comm.save(comm_p)

    state = BidState(
        run_id="run-1", body_md_path=str(body / "body.md"),
        forms_docx_path=str(forms_p), commercial_docx_path=str(comm_p),
    )
    updates = asm.assemble_node(state)
    doc = Document(updates["draft_docx_path"])
    texts = [p.text for p in doc.paragraphs]
    assert "总体思路内容。" in texts                   # 正文尾部追加
    assert "业绩：合同一份" in texts                   # commercial 兜底追加
    assert texts.count("封面页") == 1                  # 壳文本去重


def test_assemble_version_increments(tmp_path: Path, monkeypatch):
    state = _state(tmp_path, monkeypatch, version=1, forms=True, commercial=False, deviation=False)
    updates = asm.assemble_node(state)
    assert Path(updates["draft_docx_path"]).name == "标书草稿_v2.docx"
    assert (run_dir(state) / "07_draft" / "latest.txt").read_text(encoding="utf-8") == "2"
