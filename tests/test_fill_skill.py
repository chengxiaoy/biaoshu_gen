"""fill_skill 单测：以带下划线填空/下划线字符/无空白的合成模板验证填写语义。"""
from pathlib import Path

from docx import Document

from biaoshu_gen.fill_skill import (
    dump_fill_points, fill_all_blanks, fill_blank, fill_cell, find_para,
    insert_picture_after, replace_in_para, run_fill_plan,
)

# 1x1 透明 PNG（构造插图用，无需 PIL）
_PNG1 = (b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
         b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
         b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82")


def _make_template(path: Path) -> Path:
    d = Document()
    p = d.add_paragraph()
    p.add_run("项目名称：")
    blank = p.add_run("        ")
    blank.underline = True                       # (a) 带下划线的空白 run
    p2 = d.add_paragraph()
    p2.add_run("投标人（签章）：")
    p2.add_run("＿＿＿＿＿＿")                   # (b) 下划线字符 run
    p3 = d.add_paragraph()
    p3.add_run("日期：")                         # (c) 无空白无下划线
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "名称"
    d.save(path)
    return path


def _make_boundary_template(path: Path) -> Path:
    """含 投标人：/投标人（签章）：/投标人地址：/日期：无 四段，覆盖标签边界与空位两种误填。"""
    d = Document()
    p = d.add_paragraph()
    p.add_run("投标人：")
    b = p.add_run("        ")
    b.underline = True                       # 空白 run
    p2 = d.add_paragraph()
    p2.add_run("投标人（签章）：")
    p2.add_run("＿＿＿＿＿＿")                # 下划线字符 run
    p3 = d.add_paragraph()
    p3.add_run("投标人地址：")
    b3 = p3.add_run("        ")
    b3.underline = True
    p4 = d.add_paragraph()
    p4.add_run("日期：无")                    # 无空位
    d.save(path)
    return path


def test_fill_all_blanks_stops_at_label_boundary(tmp_path: Path):
    """前缀匹配须停在标签边界：投标人 命中 投标人：/投标人（签章）：，不得误中 投标人地址。"""
    d = Document(str(_make_boundary_template(tmp_path / "t.docx")))
    n = fill_all_blanks(d, "投标人", "测试公司")
    assert n == 2
    ps = [p.text for p in d.paragraphs]
    assert any(p == "投标人：测试公司" for p in ps)
    assert any(p.startswith("投标人（签章）：测试公司") for p in ps)
    addr = next(p for p in ps if p.startswith("投标人地址"))
    assert "测试公司" not in addr            # 地址段未被触碰


def test_fill_all_blanks_does_not_inject_into_slotless_para(tmp_path: Path):
    """无填空位的段落不硬插值（防 '投标人地址：无' 被塞入公司名）。"""
    d = Document(str(_make_boundary_template(tmp_path / "t.docx")))
    n = fill_all_blanks(d, "日期", "2026-08-20")
    assert n == 0
    p = next(p for p in d.paragraphs if p.text.startswith("日期"))
    assert p.text == "日期：无"              # 原文未变，未插入 run


def test_fill_blank_on_underlined_blank_run(tmp_path: Path):
    d = Document(str(_make_template(tmp_path / "t.docx")))
    fill_blank(d, "项目名称：", "演示项目")
    p = find_para(d, "项目名称：")
    assert p.text == "项目名称：演示项目"
    assert any(r.text == "演示项目" and r.underline for r in p.runs)   # 值落在线上且保留下划线


def test_fill_blank_on_underscore_run(tmp_path: Path):
    d = Document(str(_make_template(tmp_path / "t.docx")))
    fill_blank(d, "投标人（签章）：", "测试公司")
    p = find_para(d, "投标人（签章）：")
    assert p.text == "投标人（签章）：测试公司＿＿"                     # 替换线内并留余线，不附加线后


def test_fill_blank_inserts_underlined_run_when_no_blank(tmp_path: Path):
    d = Document(str(_make_template(tmp_path / "t.docx")))
    fill_blank(d, "日期：", "2026-08-20")
    p = find_para(d, "日期：")
    assert "2026-08-20" in p.text
    assert any("2026-08-20" in r.text and r.underline for r in p.runs)  # 插入的 run 自带下划线


def test_replace_and_cell_and_picture(tmp_path: Path):
    path = _make_template(tmp_path / "t.docx")
    img = tmp_path / "lic.png"
    img.write_bytes(_PNG1)
    d = Document(str(path))
    replace_in_para(d, "日期：", "日期", "签署日期")
    assert find_para(d, "签署日期：").text.startswith("签署日期")
    fill_cell(d, 0, 1, 1, "1 套")
    assert d.tables[0].rows[1].cells[1].paragraphs[0].text == "1 套"
    insert_picture_after(d, "项目名称：", str(img), caption="附：证照")
    assert len(d.inline_shapes) == 1
    assert any("附：证照" in p.text for p in d.paragraphs)


def test_run_fill_plan_batch_and_errors(tmp_path: Path):
    """声明式清单：一次执行多种 op；单条失败收集错误不中断。"""
    path = _make_template(tmp_path / "t.docx")
    img = tmp_path / "lic.png"
    img.write_bytes(_PNG1)
    plan = [
        {"op": "blank", "prefix": "项目名称：", "value": "演示项目"},
        {"op": "cell", "table_header": ["名称"], "row": 1, "col": 1, "value": "1 套"},
        {"op": "picture", "prefix": "项目名称：", "img": str(img)},
        {"op": "blank", "prefix": "不存在的段落：", "value": "X"},      # 应报错不中断
    ]
    out = tmp_path / "out.docx"
    errors = run_fill_plan(str(path), str(out), plan)
    d = Document(str(out))
    assert find_para(d, "项目名称：").text == "项目名称：演示项目"      # 前三条已生效
    assert len(d.inline_shapes) == 1
    assert len(errors) == 1 and "不存在的段落" in errors[0]


def test_dump_fill_points(tmp_path: Path):
    path = _make_template(tmp_path / "t.docx")
    d = Document(str(path))
    text = dump_fill_points(d)
    assert "项目名称" in text and "[T0]" in text and "名称" in text
