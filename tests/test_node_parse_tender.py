import json
from pathlib import Path

from docx import Document
from pydantic_ai import Agent
from pydantic_ai.messages import ModelResponse, ToolCallPart, UserPromptPart
from pydantic_ai.models.function import AgentInfo, FunctionModel

from biaoshu_gen.docx_io import DocxSection
from biaoshu_gen.nodes import DEFAULT_NODES, NODE_NAMES
from biaoshu_gen.nodes import parse_tender as pt
from biaoshu_gen.schemas import (
    InvalidationItems, ScoringStandards, TenderMetadata, TenderRequirements,
)
from biaoshu_gen.state import BidState, run_dir


def _last_user_content(messages) -> str:
    last = messages[-1]
    content = getattr(last, "content", None)
    if content is None:
        content = next(
            (p.content for p in getattr(last, "parts", []) if isinstance(p, UserPromptPart)),
            "",
        )
    return str(content)


def _state(tmp_path: Path, monkeypatch) -> BidState:
    monkeypatch.chdir(tmp_path)   # data_dir 相对路径 -> tmp
    tender = tmp_path / "tender.docx"
    d = Document()
    d.add_heading("第一章 招标公告", level=1)
    d.add_paragraph("项目名称：演示项目")
    d.add_heading("第二章 技术要求", level=1)
    d.add_paragraph("系统需支持 1000 并发。")
    d.add_heading("第三章 评标办法", level=1)
    d.add_paragraph("价格分：最低价得 100 分。")
    d.save(tender)
    return BidState(run_id="run-1", tender_path=str(tender))


def test_node_names_registry():
    assert NODE_NAMES[0] == "parse_tender" and len(NODE_NAMES) == 12
    assert DEFAULT_NODES["parse_tender"] is pt.parse_tender_node
    assert callable(DEFAULT_NODES["extract_template"])  # 未实现 -> stub


def test_classify_sections_keyword_routing():
    """纯代码路由：标题命中 + 上级章节继承 + 无命中不入组。"""
    secs = [
        DocxSection(0, "(前言)", "抬头"),
        DocxSection(1, "第五章 评标办法", "x"),
        DocxSection(3, "比较和评价", "y"),          # 无关键词，继承上级"评标"
        DocxSection(3, "应予废标的情形", "z"),       # 自身命中废标 -> invalidation+scoring
        DocxSection(1, "第七章 投标文件格式", "w"),
        DocxSection(4, "投标函", "v"),               # 与其上级均无命中
    ]
    r = pt.classify_sections(secs)
    assert 2 in r["scoring"] and 3 in r["scoring"]      # 继承
    assert 4 in r["scoring"] and 4 in r["invalidation"]
    assert 5 not in r["scoring"] and 6 not in r["scoring"]
    assert 1 not in r["metadata"]                       # 前言不误入


def test_parse_tender_routes_sections_by_keywords(tmp_path: Path, monkeypatch):
    state = _state(tmp_path, monkeypatch)
    captured: list[tuple[type, str]] = []
    presets: dict[type, dict] = {
        TenderMetadata: {"project_name": "演示项目"},
        TenderRequirements: {"tech_requirements": ["1000 并发"]},
        ScoringStandards: {"price_rules": "最低价得 100 分"},
    }

    def make(output_type, system_prompt, retries=2):
        async def fn(messages, info: AgentInfo):
            captured.append((output_type, _last_user_content(messages)))
            tool_name = info.output_tools[0].name if info.output_tools else "final_result"
            out = presets.get(output_type, {"items": []})
            return ModelResponse(parts=[ToolCallPart(tool_name=tool_name, args=json.dumps(out))])
        return Agent(model=FunctionModel(fn), output_type=output_type,
                     system_prompt=system_prompt, retries=retries)

    monkeypatch.setattr(pt, "make_agent", make)
    updates = pt.parse_tender_node(state)

    d = run_dir(state) / "01_parse"
    assert (d / "tender.md").exists()
    assert (d / "metadata.yaml").exists() and (d / "scoring.yaml").exists()
    assert (d / "routing.yaml").exists()               # 路由透明化
    assert updates["metadata"].project_name == "演示项目"
    assert updates["requirements"].tech_requirements == ["1000 并发"]

    # 分节路由断言：每组抽取只看到本组章节内容（关键词路由，无 LLM 分类调用）
    called_types = {t for t, _ in captured}
    assert called_types == {TenderMetadata, TenderRequirements, ScoringStandards}
    meta_prompt = next(p for t, p in captured if t is TenderMetadata)
    assert "项目名称：演示项目" in meta_prompt and "评标办法" not in meta_prompt
    scoring_prompt = next(p for t, p in captured if t is ScoringStandards)
    assert "最低价得 100 分" in scoring_prompt and "技术要求" not in scoring_prompt


def test_parse_tender_keyword_sections_routed_to_invalidation(tmp_path: Path, monkeypatch):
    """标题含 废标/无效/扣分/偏离 的章节自动进入 invalidation 组。"""
    monkeypatch.chdir(tmp_path)
    tender = tmp_path / "t2.docx"
    d = Document()
    d.add_heading("废标条款", level=1)
    d.add_paragraph("逾期送达的投标文件将被拒收。")
    d.save(tender)
    state = BidState(run_id="run-2", tender_path=str(tender))

    presets: dict[type, dict] = {
        InvalidationItems: {"items": [{"kind": "废标项", "requirement": "不得逾期送达"}]},
    }

    def make(output_type, system_prompt, retries=2):
        async def fn(messages, info: AgentInfo):
            tool_name = info.output_tools[0].name if info.output_tools else "final_result"
            out = presets.get(output_type, {})
            return ModelResponse(parts=[ToolCallPart(tool_name=tool_name, args=json.dumps(out))])
        return Agent(model=FunctionModel(fn), output_type=output_type,
                     system_prompt=system_prompt, retries=retries)

    monkeypatch.setattr(pt, "make_agent", make)
    updates = pt.parse_tender_node(state)
    assert updates["invalidation"].items[0].kind == "废标项"   # 关键词路由使抽取确实发生
