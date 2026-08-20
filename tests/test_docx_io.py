from pathlib import Path

from docx import Document

from biaoshu_gen.docx_io import (
    DocxSection, copy_docx, docx_to_markdown, docx_to_sections, markdown_to_docx,
)


def _make_tender_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("第一章 招标公告", level=1)
    doc.add_paragraph("项目名称：测试项目")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "名称"
    t.cell(0, 1).text = "数量"
    t.cell(1, 0).text = "应用软件 A"
    t.cell(1, 1).text = "1 套"
    doc.save(path)


def test_docx_to_markdown_keeps_order_and_table(tmp_path: Path):
    p = tmp_path / "t.docx"
    _make_tender_docx(p)
    md = docx_to_markdown(p)
    assert "# 第一章 招标公告" in md
    assert "项目名称：测试项目" in md
    assert "| 名称 | 数量 |" in md
    assert "| 应用软件 A | 1 套 |" in md
    assert md.index("招标公告") < md.index("项目名称") < md.index("应用软件 A")


def test_docx_to_sections_splits_by_heading(tmp_path: Path):
    p = tmp_path / "t2.docx"
    doc = Document()
    doc.add_paragraph("抬头说明")                      # 首个标题前 → 前言
    doc.add_heading("第一章 招标公告", level=1)
    doc.add_paragraph("项目名称：测试项目")
    doc.add_heading("评标办法", level=1)
    doc.add_paragraph("最低价得 100 分")
    doc.save(p)
    secs = docx_to_sections(p)
    assert [(s.level, s.title) for s in secs] == [
        (0, "(前言)"), (1, "第一章 招标公告"), (1, "评标办法")]
    assert "抬头说明" in secs[0].content
    assert "项目名称" in secs[1].content
    assert "100 分" in secs[2].content


def test_markdown_to_docx_headings_and_list(tmp_path: Path):
    doc = Document()
    markdown_to_docx(doc, "# 总体方案\n\n本章说明总体设计。\n\n- 要点一\n- 要点二\n")
    paras = [p.text for p in doc.paragraphs]
    styles = [p.style.name for p in doc.paragraphs]
    assert "总体方案" in paras and "本章说明总体设计。" in paras and "要点一" in paras
    assert any("Heading 1" in s for s in styles)
    assert any("List" in s for s in styles)




def test_copy_docx(tmp_path: Path):
    src = tmp_path / "tpl.docx"
    _make_tender_docx(src)
    doc = copy_docx(src, tmp_path / "tpl_copy.docx")
    assert "第一章 招标公告" in "\n".join(p.text for p in doc.paragraphs)
    assert (tmp_path / "tpl_copy.docx").exists()


def test_markdown_to_docx_style_fallback():
    """中文模板底稿常缺 List Bullet 等样式：样式缺失时回退普通段落，不崩溃。"""
    doc = Document()
    for name in ("List Bullet", "List Number", "Heading 2"):
        el = doc.styles[name].element
        el.getparent().remove(el)
    markdown_to_docx(doc, "## 某标题\n\n- 要点一\n\n1. 步骤一\n")
    texts = [p.text for p in doc.paragraphs]
    assert "要点一" in texts and "步骤一" in texts and "某标题" in texts


def test_template_has_section(tmp_path: Path):
    from biaoshu_gen.docx_io import template_has_section

    no = tmp_path / "no.docx"
    d = Document()
    d.add_paragraph("投标函")
    d.save(no)
    assert template_has_section(no, "偏离") is False

    yes = tmp_path / "yes.docx"
    d = Document()
    d.add_paragraph("偏离表")
    d.save(yes)
    assert template_has_section(yes, "偏离") is True
    assert template_has_section(tmp_path / "missing.docx", "偏离") is False
