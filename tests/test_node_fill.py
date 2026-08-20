from pathlib import Path

from biaoshu_gen.harness import prepare_agent_workspace
from biaoshu_gen.kb import KnowledgeBase
from biaoshu_gen.nodes import commercial as com
from biaoshu_gen.nodes import deviation_table as dev
from biaoshu_gen.nodes import fill_forms as ff
from biaoshu_gen.state import BidState, run_dir


def _fake_run(captured):
    def fake(task):
        captured.append((task.cwd, task.prompt, task.expected_outputs))
        for p in task.expected_outputs:
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_bytes(b"fake-docx")
        return task.expected_outputs
    return fake


def _patch_fill_harness(monkeypatch, captured, mod=None):
    """fill 节点经 fill_context.run_fill_node 调用 harness，故 patch 该模块。"""
    from biaoshu_gen import fill_context
    monkeypatch.setattr(fill_context, "run_harness_task", _fake_run(captured))


def _base_state(tmp_path: Path, monkeypatch) -> BidState:
    monkeypatch.chdir(tmp_path)
    state = BidState(run_id="run-1", kb_dir=str(tmp_path / "kb"))
    (tmp_path / "kb").mkdir(exist_ok=True)
    (tmp_path / "kb" / "简介.md").write_text("公司具备 CMMI5。", encoding="utf-8")
    (tmp_path / "kb" / "营业执照.jpg").write_bytes(b"\xff\xd8img")
    parse = run_dir(state) / "01_parse"
    parse.mkdir(parents=True)
    (parse / "tender.md").write_text("# 招标公告", encoding="utf-8")
    (parse / "invalidation.yaml").write_text("items: []\n", encoding="utf-8")
    (parse / "metadata.yaml").write_text("project_name: 演示\n", encoding="utf-8")
    (parse / "requirements.yaml").write_text("tech_requirements: []\n", encoding="utf-8")
    (parse / "scoring.yaml").write_text("technical_rules: []\n", encoding="utf-8")
    (run_dir(state) / "03_facts.yaml").write_text("schedule: 90 天\n", encoding="utf-8")
    return state


def test_dump_summary_contains_text_and_images(tmp_path: Path):
    kb_dir = tmp_path / "kb"
    kb_dir.mkdir()
    (kb_dir / "a.md").write_text("具备 ISO27001。", encoding="utf-8")
    (kb_dir / "lic.jpg").write_bytes(b"\xff\xd8x")
    out = KnowledgeBase.load(kb_dir).dump_summary(tmp_path / "kb.md")
    text = out.read_text(encoding="utf-8")
    assert "ISO27001" in text and str((kb_dir / "lic.jpg").resolve()) in text


def _with_template(tmp_path: Path, monkeypatch, text: str = "偏离表") -> BidState:
    """构造含指定段落（默认偏离表）的响应模板（判定依据为响应模板）。"""
    from docx import Document
    state = _base_state(tmp_path, monkeypatch)
    tpl = tmp_path / "标书模板.docx"
    d = Document()
    d.add_paragraph(text)
    d.save(tpl)
    return state.model_copy(update={"template_docx_path": str(tpl)})


def test_three_fill_nodes_isolated_workspaces(tmp_path: Path, monkeypatch):
    state = _with_template(tmp_path, monkeypatch, text="商务部分\n偏离表")
    captured = []
    _patch_fill_harness(monkeypatch, captured)
    u1 = ff.fill_forms_node(state)
    u2 = dev.deviation_table_node(state)
    u3 = com.commercial_node(state)

    assert u1["forms_docx_path"].endswith(str(Path("06_fill/forms/forms.docx")))
    assert u2["deviation_docx_path"].endswith("deviation.docx")
    assert u3["commercial_docx_path"].endswith("commercial.docx")
    assert len({c[0] for c in captured}) == 3            # forms + deviation + commercial 工作区隔离
    # 标准工作区内容：tender.md / invalidation.yaml / kb.md
    ws = run_dir(state) / "06_fill" / "forms"
    assert (ws / "tender.md").exists() and (ws / "kb.md").exists()
    assert "CMMI5" in (ws / "kb.md").read_text(encoding="utf-8")
    # 各节点附加输入正确
    assert (ws / "metadata.yaml").exists() and (ws / "facts.yaml").exists()
    assert (run_dir(state) / "06_fill" / "commercial" / "scoring.yaml").exists()


def test_deviation_skipped_without_template(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    state = _base_state(tmp_path, monkeypatch)
    assert dev.deviation_table_node(state) == {"deviation_docx_path": ""}


def test_commercial_skipped_without_template(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    state = _base_state(tmp_path, monkeypatch)
    assert com.commercial_node(state) == {"commercial_docx_path": ""}


def test_commercial_skipped_when_template_no_commercial(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    state = _with_template(tmp_path, monkeypatch, text="偏离表")
    assert com.commercial_node(state) == {"commercial_docx_path": ""}


def test_commercial_fills_template_with_commercial(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    state = _with_template(tmp_path, monkeypatch, text="商务部分\n业绩证明文件")
    captured = []
    _patch_fill_harness(monkeypatch, captured)
    updates = com.commercial_node(state)
    assert updates["commercial_docx_path"].endswith("commercial.docx")
    assert len(captured) == 1
    assert "商务部分" in captured[0][1]                  # prompt 强调按模板商务部分填写
    assert "不得删减" not in captured[0][1] or "标书模板.docx" in captured[0][1]


def test_prepare_agent_workspace_base_inputs(tmp_path: Path, monkeypatch):
    state = _base_state(tmp_path, monkeypatch)
    tpl = tmp_path / "标书模板.docx"
    tpl.write_bytes(b"tpl")
    state = state.model_copy(update={"template_docx_path": str(tpl)})
    (tmp_path / "extra.yaml").write_text("e: 1", encoding="utf-8")
    ws = prepare_agent_workspace(state, "06_fill/forms", [
        (tmp_path / "extra.yaml", "extra.yaml")])
    for name in ("tender.md", "invalidation.yaml", "标书模板.docx", "kb.md", "extra.yaml",
                 "fill_skill.py"):
        assert (ws / name).exists(), name


def test_fill_prompts_preinject_context(tmp_path: Path, monkeypatch):
    """模板地图/facts/图片路径预注入 prompt，harness 无需读文件探查。"""
    monkeypatch.chdir(tmp_path)
    from docx import Document
    state = _base_state(tmp_path, monkeypatch)
    tpl = tmp_path / "标书模板.docx"
    d = Document()
    d.add_paragraph("项目名称：＿＿＿")
    d.add_paragraph("商务部分")
    d.save(tpl)
    state = state.model_copy(update={"template_docx_path": str(tpl)})

    captured = []
    _patch_fill_harness(monkeypatch, captured)
    com.commercial_node(state)
    prompt = captured[0][1]
    assert "模板可填点地图" in prompt and "项目名称" in prompt   # 地图已注入
    assert "facts.yaml 全文" in prompt and "90 天" in prompt    # facts 已注入
    assert "营业执照.jpg" in prompt                              # 图片绝对路径已注入


def test_prefill_known_fills_deterministic_values(tmp_path: Path, monkeypatch):
    """确定值（项目名称/编号/投标人等）由代码预填进模板副本，prompt 标注勿重复。"""
    monkeypatch.chdir(tmp_path)
    from docx import Document
    state = _base_state(tmp_path, monkeypatch)
    tpl = tmp_path / "标书模板.docx"
    d = Document()
    for label in ("项目名称：", "项目编号：", "投标人（签章）：", "法定代表人：", "采购人名称："):
        p = d.add_paragraph()
        p.add_run(label)
        blank = p.add_run("        ")
        blank.underline = True
    d.add_paragraph("商务部分")
    d.save(tpl)
    state = state.model_copy(update={"template_docx_path": str(tpl)})
    # facts 带 template_fields 与企业资料
    from biaoshu_gen.business import ensure_business_fields
    ensure_business_fields(state)
    fp = run_dir(state) / "03_facts.yaml"
    fp.write_text(fp.read_text(encoding="utf-8")
                  + 'template_fields:\n  项目名称: 演示项目\n  项目编号: DEMO-001\n  采购人名称: 某某局\n',
                  encoding="utf-8")

    from biaoshu_gen.fill_context import prefill_known
    from docx import Document as D
    doc = D(str(tpl))
    summary = prefill_known(doc, state)
    doc.save(tpl)
    d2 = D(str(tpl))
    texts = [p.text for p in d2.paragraphs]
    assert "项目名称：演示项目" in texts
    assert "项目编号：DEMO-001" in texts
    assert "采购人名称：某某局" in texts
    assert any(t.startswith("投标人（签章）：某某科技") for t in texts)   # mock 企业名已预填
    assert any(t.startswith("法定代表人：法定代表人") for t in texts)
    assert "项目名称×1" in summary and "投标人×1" in summary
