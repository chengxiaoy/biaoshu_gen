"""标书模板填写 skill：表格填写 / 下划线填空 / 插入图片 的可复用原语。

设计要点（源自 fill 阶段 harness 实战脚本的提炼）：
- **前缀锚定**而非魔法下标：以段落文本前缀定位（如 "项目名称："），模板微调不致错位；
- **下划线填空**：值填*在下划线上*（优先填带下划线格式的空白 run；其次替换下划线字符 run
  并保留少量余线；再无则复制邻近格式插入带下划线的 run）——修复"值附加在下划线之后"的问题；
- **图片**：WEBP 伪装 .jpg 时机械转码 PNG（不读取内容）；插图带居中与可选图注。

供 fill 阶段三个 harness 节点直接 import 使用（工作区内会自动放置本文件副本）。
"""
import copy as _copy
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

UNDERLINE_CHARS = set("＿＿___―—-") - set("")  # 全角/半角下划线

# 标签之后的合法边界：分隔符/括号/空白/段末/下划线字符（防 "投标人" 误中 "投标人地址"）
_BOUNDARY_CHARS = set("：:（）() \t") | UNDERLINE_CHARS


def _has_fill_slot(p: Paragraph) -> bool:
    """段落是否真实存在下划线填空位（下划线空白 run 或下划线字符 run）；无则跳过不硬插。"""
    if any(r.text and not r.text.strip() and _is_underlined(r) for r in p.runs):
        return True
    return any((r.text or "").strip() and set((r.text or "").strip()) <= UNDERLINE_CHARS
               for r in p.runs)


def _is_underlined(run) -> bool:
    rPr = run._element.rPr
    if rPr is None:
        return False
    u = rPr.find(qn("w:u"))
    return u is not None and u.get(qn("w:val")) == "single"


def find_para(doc, prefix: str) -> Paragraph:
    """按文本前缀定位段落；找不到抛 RuntimeError（带提示便于 harness 自纠）。"""
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise RuntimeError(f"找不到以 {prefix!r} 开头的段落；请核对模板文本")


def _fill_blank_in_para(p: Paragraph, value: str) -> None:
    """在单个段落的填空线上填 value（保留下划线格式，不追加到线后）。"""
    blanks = [r for r in p.runs if r.text and not r.text.strip()]
    underlined = [r for r in blanks if _is_underlined(r)]
    if underlined:
        underlined[0].text = value
        for r in blanks:                       # 其余空白 run 清空（避免重复落值）
            if r is not underlined[0]:
                r.text = ""
        return
    for r in p.runs:                            # 下划线字符 run（＿＿＿/___）
        t = (r.text or "").strip()
        if t and set(t) <= UNDERLINE_CHARS:
            r.text = f"{value}{'＿' * 2}"       # 值落在线上并保留余线
            return
    # 无空白也无下划线字符：复制末 run 格式插入带下划线的值 run（位置在段内，非段后附加）
    new_r = p.add_run(f" {value} ")
    last = p.runs[-2] if len(p.runs) >= 2 else p.runs[0]   # 复制标签 run 格式
    if last is not None and last._element.rPr is not None:
        new_r._element.insert(0, _copy.deepcopy(last._element.rPr))
    rPr = new_r._element.get_or_add_rPr()
    u = rPr.find(qn("w:u"))
    if u is None:
        u = rPr.makeelement(qn("w:u"), {})
        rPr.append(u)
    u.set(qn("w:val"), "single")


def fill_blank(doc, prefix: str, value: str) -> Paragraph:
    """在 prefix 段落的填空线上填 value（首个匹配段落）。"""
    p = find_para(doc, prefix)
    _fill_blank_in_para(p, value)
    return p


def fill_all_blanks(doc, prefix: str, value: str) -> int:
    """把**所有**以 prefix 开头的段落的填空线都填上 value，返回填写段数（预填已知值用）。

    两道护栏（区别于 blank op 的 fill_blank，预填宁可少填也不可错填）：
    - 标签边界：prefix 之后须是分隔符/括号/空白/段末，避免 "投标人" 误中 "投标人地址"；
    - 空位门槛：段落须真实存在下划线填空位，无空位段落不硬插值（防 "投标人地址：无" 被塞值）。
    """
    n = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text.startswith(prefix):
            continue
        rest = text[len(prefix):]
        if rest and rest[0] not in _BOUNDARY_CHARS:
            continue
        if not _has_fill_slot(p):
            continue
        _fill_blank_in_para(p, value)
        n += 1
    return n


def replace_in_para(doc, prefix: str, old: str, new: str) -> Paragraph:
    """段内文本替换（保留首个 run 格式，整段合并）。"""
    p = find_para(doc, prefix)
    full = "".join(r.text for r in p.runs)
    if old not in full:
        raise RuntimeError(f"{prefix!r} 段落中未找到 {old!r}：{full[:60]!r}")
    p.runs[0].text = full.replace(old, new)
    for r in p.runs[1:]:
        r.text = ""
    return p


def fill_cell(doc, table_idx: int, row: int, col: int, text: str):
    """填表格单元格（保留表格结构）。"""
    para = doc.tables[table_idx].rows[row].cells[col].paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)
    return para


def ensure_readable_img(path: str) -> str:
    """WEBP 伪装 .jpg 的文件 python-docx 无法嵌入：机械转码 PNG（不读取内容）。"""
    with open(path, "rb") as fh:
        head = fh.read(12)
    if head.startswith(b"RIFF"):
        from PIL import Image
        im = Image.open(path)
        tmp = os.path.join(os.path.dirname(os.path.abspath(path)),
                           "_conv_" + os.path.splitext(os.path.basename(path))[0] + ".png")
        im.save(tmp, "PNG")
        return tmp
    return path


def _new_para_after(p: Paragraph) -> Paragraph:
    new_p = _copy.deepcopy(p._p)
    p._p.addnext(new_p)
    np_ = Paragraph(new_p, p._parent)
    for r in np_.runs:
        r.text = ""
    return np_


def insert_picture_after(doc, prefix: str, img: str, width_inch: float = 5.6,
                         caption: str | None = None) -> Paragraph:
    """在 prefix 段落之后插入居中图片（可选图注），返回可继续链式插入的锚段。"""
    p = find_para(doc, prefix)
    np_ = _new_para_after(p)
    np_.alignment = 1                       # center
    r = np_.add_run()
    r.add_picture(ensure_readable_img(img), width=Inches(width_inch))
    last = np_
    if caption:
        cp = _new_para_after(np_)
        cr = cp.add_run(caption)
        cr.font.size = Pt(9)
        cp.alignment = 1
        last = cp
    return last


# ---------------- 声明式填空清单（一次执行、批量报错，压缩 harness 轮次） ----------------

def find_table(doc, *header_keywords: str) -> int:
    """按表头关键词定位表格（表头行含全部关键词），返回下标；找不到抛 RuntimeError。"""
    for i, t in enumerate(doc.tables):
        head = " ".join(c.text for c in t.rows[0].cells)
        if all(k in head for k in header_keywords):
            return i
    raise RuntimeError(f"找不到表头含 {header_keywords} 的表格")


def dump_fill_points(doc) -> str:
    """一次性输出模板全部可填点地图：段落（下标/文本/是否含填空线）+ 表格（下标/表头）。"""
    lines = ["== 段落 =="]
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        has_blank = any(r.text and not r.text.strip() and _is_underlined(r) for r in p.runs) \
            or any((r.text or "").strip() and set((r.text or "").strip()) <= UNDERLINE_CHARS
                   for r in p.runs)
        lines.append(f"[{i}]{'(线)' if has_blank else ''} {t[:50]}")
    lines.append("== 表格 ==")
    for i, t in enumerate(doc.tables):
        head = " | ".join(c.text.strip()[:8] for c in t.rows[0].cells)
        lines.append(f"[T{i}] {head}  ({len(t.rows)}行)")
    return "\n".join(lines)


def run_fill_plan(template: str, output: str, plan: list[dict]) -> list[str]:
    """按填空清单一次性执行全部操作；单条失败不中断，返回错误清单供批量修正。

    plan 条目（op 必填）：
      {"op":"blank","prefix":"项目名称：","value":"X"}                 # 下划线填空
      {"op":"replace","prefix":"致：","old":"（采购人）","new":"X"}      # 段内替换
      {"op":"cell","table":0,"row":1,"col":2,"value":"X"}              # 按下标填格
      {"op":"cell","table_header":["序号","名称"],"row":1,"col":1,...} # 按表头定位填格
      {"op":"picture","prefix":"备注：","img":"C:/...jpg","width":4.8,"caption":"附：X"}
      {"op":"append","prefix":"投标人名称：","value":"X"}               # 段末追加（无填空线时）
    """
    errors: list[str] = []
    doc = Document(template)
    for i, op in enumerate(plan):
        try:
            kind = op["op"]
            if kind == "blank":
                fill_blank(doc, op["prefix"], op["value"])
            elif kind == "replace":
                replace_in_para(doc, op["prefix"], op["old"], op["new"])
            elif kind == "cell":
                t = op.get("table")
                if t is None:
                    t = find_table(doc, *op["table_header"])
                fill_cell(doc, int(t), int(op["row"]), int(op["col"]), op["value"])
            elif kind == "picture":
                insert_picture_after(doc, op["prefix"], op["img"],
                                     float(op.get("width", 5.6)), op.get("caption"))
            elif kind == "append":
                find_para(doc, op["prefix"]).add_run(op["value"])
            else:
                raise RuntimeError(f"未知 op: {kind}")
        except Exception as e:              # 收集错误继续执行，供一次修正
            errors.append(f"[{i}] {op.get('op')} {op.get('prefix', op.get('table_header', ''))}: {e}")
    doc.save(output)
    return errors
