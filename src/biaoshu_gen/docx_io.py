"""docx 与 Markdown 的双向转换、模板复制、文档合并。"""
import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(doc: DocumentType):
    """按文档真实顺序产出段落与表格。"""
    from docx.oxml.ns import qn
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _table_md(table: Table) -> str:
    lines = []
    for row in table.rows:
        cells = [c.text.replace("\n", " ").replace("|", "/").strip() for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


@dataclass
class DocxSection:
    """docx 按标题切出的章节。level=0 表示首个标题前的前言。"""
    level: int          # 0=前言, 1~4=Heading N
    title: str
    content: str        # 本节正文 Markdown（含表格）


_HEADING_RE = re.compile(r"(?:heading|标题)\s*(\d)", re.IGNORECASE)


def docx_to_sections(path: Path) -> list[DocxSection]:
    doc = Document(str(path))
    sections: list[DocxSection] = []
    cur: DocxSection | None = None

    def flush(text: str) -> None:
        nonlocal cur
        if cur is None:
            cur = DocxSection(0, "(前言)", "")
        if text:
            cur.content = (cur.content + "\n\n" + text).strip()

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            m = _HEADING_RE.match((block.style.name or "").strip())
            if m and text:
                if cur is not None:
                    sections.append(cur)
                cur = DocxSection(int(m.group(1)), text, "")
            else:
                flush(text)
        else:
            flush(_table_md(block))
    if cur is not None:
        sections.append(cur)
    return sections


def docx_to_markdown(path: Path) -> str:
    parts: list[str] = []
    for s in docx_to_sections(path):
        if s.level:
            parts.append("#" * s.level + " " + s.title)
        if s.content:
            parts.append(s.content)
    return "\n\n".join(parts) + "\n"


def _add_styled(doc: DocumentType, text: str, style: str):
    """按样式加段落；样式缺失（常见于中文模板底稿）时回退为普通段落。"""
    try:
        return doc.add_paragraph(text, style=style)
    except KeyError:
        return doc.add_paragraph(text)


def markdown_to_docx(doc: DocumentType, md: str) -> None:
    """极量版 Markdown → docx：标题/列表/段落（POC 够用）。"""
    for line in md.splitlines():
        s = line.strip()
        if not s:
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            try:
                doc.add_heading(m.group(2), level=len(m.group(1)))
            except KeyError:                      # 模板缺 Heading N 样式时回退
                doc.add_paragraph(m.group(2))
        elif s.startswith(("- ", "* ")):
            _add_styled(doc, s[2:], "List Bullet")
        elif re.match(r"^\d+\.\s+", s):
            _add_styled(doc, re.sub(r"^\d+\.\s+", "", s), "List Number")
        else:
            doc.add_paragraph(re.sub(r"\*\*(.+?)\*\*", r"\1", s))


def append_docx(dest: DocumentType, src_path: Path) -> None:
    """把 src 文档 body 的段落/表格深拷贝追加到 dest（跨文档移动需要 deepcopy）。"""
    import copy as _copy

    src = Document(str(src_path))
    sect_pr = dest.element.body.sectPr
    for child in src.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            el = _copy.deepcopy(child)
            if sect_pr is not None:
                sect_pr.addprevious(el)   # sectPr 必须是 body 最后一个子元素（ECMA-376）
            else:
                dest.element.body.append(el)


def copy_docx(src: Path, dest: Path) -> DocumentType:
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(src, dest)
    return Document(str(dest))


@dataclass
class DocxBlockRange:
    """标题区间：锚标题 + 其管辖的 body 子元素（到下一同级/更高级标题前，不含 sectPr）。"""
    title: str
    level: int
    elements: list


def docx_block_ranges(doc: DocumentType) -> list[DocxBlockRange]:
    """按标题把文档切成区间（assemble 分段定位/替换的基础）。"""
    ranges: list[DocxBlockRange] = []
    cur: DocxBlockRange | None = None
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "sectPr":
            continue                        # 分节属性固定在 body 尾部，不参与区间
        level = 0
        if tag == "p":
            para = Paragraph(child, doc)
            m = _HEADING_RE.match((para.style.name or "").strip())
            if m and para.text.strip():
                if cur is not None:
                    ranges.append(cur)
                cur = DocxBlockRange(para.text.strip(), int(m.group(1)), [child])
                continue
        if cur is None:
            cur = DocxBlockRange("(前言)", 0, [])
        cur.elements.append(child)
    if cur is not None:
        ranges.append(cur)
    return ranges


def replace_elements(old_elements: list, new_elements: list) -> None:
    """就地替换：new_elements（来自其他文档，自动深拷贝）替换 old_elements，保持原位置。"""
    import copy as _copy

    if not old_elements:
        return
    anchor = old_elements[0]
    for el in new_elements:
        anchor.addprevious(_copy.deepcopy(el))
    for el in old_elements:
        el.getparent().remove(el)


def append_elements_before_sectpr(doc: DocumentType, elements: list) -> None:
    """把 elements（外部文档元素，深拷贝）追加到 body 末尾（sectPr 之前）。"""
    import copy as _copy

    sect_pr = doc.element.body.sectPr
    for el in elements:
        if sect_pr is not None:
            sect_pr.addprevious(_copy.deepcopy(el))
        else:
            doc.element.body.append(_copy.deepcopy(el))


_DEVIATION_KEYWORDS = ("偏离表", "偏离情况", "商务偏离", "技术偏离", "负偏离", "正偏离", "偏离说明")


def template_has_section(tpl_path: Path, keyword: str, table_hint: bool = False) -> bool:
    """动态判断响应模板中是否存在含 keyword 的部分（扫描段落；table_hint 时含表头）。"""
    if not tpl_path.exists():
        return False
    doc = Document(str(tpl_path))
    for p in doc.paragraphs:
        if keyword in p.text:
            return True
    if table_hint:
        for tb in doc.tables:
            for row in tb.rows[:3]:
                cells = " ".join(c.text for c in row.cells)
                if keyword in cells:
                    return True
    return False


def template_has_deviation_table(tpl_path: Path) -> bool:
    """响应模板中是否存在偏离表要求：段落关键词或含"招标…要求+响应"的表头。"""
    if not tpl_path.exists():
        return False
    if template_has_section(tpl_path, "偏离"):
        return True
    return False
