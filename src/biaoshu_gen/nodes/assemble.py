"""节点 10：拼装标书草稿 docx（纯代码，无 LLM）。

策略：**标题锚定的分段替换**——
- 底稿 = forms.docx（模板壳 + 已填的投标函/报价/资格）> 标书模板.docx > 新建；
- 技术方案正文：写入底稿"技术部分/技术方案"锚点区间（找不到锚点才追加尾部）；
- 商务部分/偏离表：在底稿与填充文档中按锚标题定位**同一区间**，整段替换（空壳 → 已填），
  规避逐块文本去重在"同节不同填充状态"下的重复拼接与误删；
- 底稿没有对应区间时，仅把填充文档中**该区间的内容**追加尾部（不整本拼接）；
- 填充文档连锚标题都没有时，兜底走整本去重追加。
"""
from pathlib import Path

from docx import Document

from ..docx_io import (
    append_elements_before_sectpr, copy_docx, docx_block_ranges,
    markdown_to_docx, replace_elements,
)
from ..state import BidState, run_dir

_TECH_KEYWORDS = ("技术方案", "技术部分", "技术标", "技术")


def _find_range(ranges, keywords: tuple[str, ...]):
    for kw in keywords:
        for r in ranges:
            if kw in r.title:
                return r
    return None


def _content_elements(md: str) -> list:
    """把 markdown 渲染到临时文档，返回其 body 元素（不含 sectPr）。"""
    scratch = Document()
    markdown_to_docx(scratch, md)
    return [el for el in scratch.element.body.iterchildren()
            if el.tag.split("}")[-1] != "sectPr"]


def _block_key(block) -> str:
    """段落取全文，表格取全部单元格文本，用于兜底去重比较。"""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if isinstance(block, Paragraph):
        return f"p:{block.text.strip()}"
    if isinstance(block, Table):
        return "t:" + "|".join(c.text.strip() for row in block.rows for c in row.cells)
    return ""


def _collect_keys(doc: Document) -> set:
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    keys: set = set()
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        block = Paragraph(child, doc) if tag == "p" else (Table(child, doc) if tag == "tbl" else None)
        if block is None:
            continue
        key = _block_key(block)
        if key:
            keys.add(key)
    return keys


def _append_docx_dedup(dest: Document, src_path: Path, existing: set) -> None:
    """兜底：整本去重追加（src 中与底稿文本相同的块跳过）。"""
    import copy as _copy

    src = Document(str(src_path))
    sect_pr = dest.element.body.sectPr
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in src.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            block = Paragraph(child, src)
        elif tag == "tbl":
            block = Table(child, src)
        else:
            continue
        key = _block_key(block)
        if not key or key in existing:
            continue
        existing.add(key)
        el = _copy.deepcopy(child)
        if sect_pr is not None:
            sect_pr.addprevious(el)
        else:
            dest.element.body.append(el)


def assemble_node(state: BidState) -> dict:
    out_dir = run_dir(state) / "07_draft"
    out_dir.mkdir(parents=True, exist_ok=True)
    version = state.draft_version + 1
    dest = out_dir / f"标书草稿_v{version}.docx"

    # 底稿：填好的 forms 优先（避免空模板壳），否则响应模板，否则新建
    if state.forms_docx_path and Path(state.forms_docx_path).exists():
        doc = copy_docx(Path(state.forms_docx_path), dest)
    elif state.template_docx_path and Path(state.template_docx_path).exists():
        doc = copy_docx(Path(state.template_docx_path), dest)
    else:
        doc = Document()
        if state.metadata and state.metadata.project_name:
            doc.add_heading(f"{state.metadata.project_name} 投标文件", level=0)

    # 技术方案正文 -> 锚定"技术部分"区间（保留锚标题，替换区间其余内容）
    body_md = Path(state.body_md_path).read_text(encoding="utf-8")
    tech = _find_range(docx_block_ranges(doc), _TECH_KEYWORDS)
    if tech is not None:
        replace_elements(tech.elements[1:], _content_elements(body_md))
    else:
        doc.add_page_break()
        markdown_to_docx(doc, "# 技术方案\n\n" + body_md)

    # 商务部分 / 偏离表 -> 同锚区间整段替换；底稿无该区间则仅追加该区间；再兜底整本去重
    for path, keywords in (
        (state.commercial_docx_path, ("商务部分", "商务")),
        (state.deviation_docx_path, ("偏离表", "偏离")),
    ):
        if not (path and Path(path).exists()):
            continue
        src = Document(str(path))
        src_range = _find_range(docx_block_ranges(src), keywords)
        base_range = _find_range(docx_block_ranges(doc), keywords)
        if src_range is not None and base_range is not None:
            replace_elements(base_range.elements, src_range.elements)
        elif src_range is not None:
            doc.add_page_break()
            append_elements_before_sectpr(doc, src_range.elements)
        else:
            doc.add_page_break()
            _append_docx_dedup(doc, Path(path), _collect_keys(doc))
    doc.save(str(dest))

    (out_dir / "latest.txt").write_text(str(version), encoding="utf-8")
    md_path = out_dir / f"标书草稿_v{version}.md"
    md_path.write_text(body_md + "\n\n（已并入填充产物：forms / deviation / commercial docx）\n",
                       encoding="utf-8")
    return {"draft_docx_path": str(dest), "draft_md_path": str(md_path), "draft_version": version}
